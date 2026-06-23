"""
candidate_ranker.py
===================
Module 1 of the Candidate Ranking System.

Pipeline
--------
1. Load candidates.jsonl line-by-line.
2. Build a rich_profile string per candidate from all meaningful fields.
3. Embed profiles with sentence-transformers (all-MiniLM-L6-v2).
4. Store embeddings in a FAISS flat-L2 index, persist index + metadata.
5. Expose get_top_k_candidates() for JD-based retrieval.

Usage
-----
    python candidate_ranker.py                  # builds index + runs demo query
    python candidate_ranker.py --rebuild        # force rebuild even if index exists

Dependencies
------------
    pip install pandas faiss-cpu sentence-transformers python-docx tqdm
"""

from __future__ import annotations

import argparse
import json
import logging
import os
import sys
import time
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

import faiss
import numpy as np
import pandas as pd
from docx import Document
from sentence_transformers import SentenceTransformer
from tqdm import tqdm

# ---------------------------------------------------------------------------
# Configuration
# ---------------------------------------------------------------------------

BASE_DIR = Path(__file__).parent.parent
CANDIDATES_JSONL = BASE_DIR / "data" / "candidates.jsonl"
SCHEMA_PATH = BASE_DIR / "data" / "candidate_schema.json"
JD_PATH = BASE_DIR / "data" / "job_description.docx"
INDEX_PATH = BASE_DIR / "outputs" / "candidates_index.faiss"
METADATA_PATH = BASE_DIR / "outputs" / "metadata.json"

EMBEDDING_MODEL = "all-MiniLM-L6-v2"
EMBEDDING_DIM = 384          # fixed for all-MiniLM-L6-v2
BATCH_SIZE = 256             # candidates embedded per batch
NOT_PROVIDED = "Not Provided"

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s | %(levelname)-8s | %(message)s",
    datefmt="%H:%M:%S",
)
logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Data classes
# ---------------------------------------------------------------------------

@dataclass
class CandidateRecord:
    """Lightweight container for a parsed candidate."""
    candidate_id: str
    rich_profile: str
    raw: dict = field(default_factory=dict, repr=False)


# ---------------------------------------------------------------------------
# Helper utilities
# ---------------------------------------------------------------------------

def _safe_str(value: Any, fallback: str = NOT_PROVIDED) -> str:
    """Return str(value) if value is truthy, else fallback."""
    if value is None or value == "" or value == [] or value == {}:
        return fallback
    return str(value).strip() or fallback


def _format_career_history(career: list[dict]) -> str:
    """Convert career history list to a readable narrative."""
    if not career:
        return NOT_PROVIDED
    parts: list[str] = []
    for job in career:
        title = _safe_str(job.get("title"))
        company = _safe_str(job.get("company"))
        industry = _safe_str(job.get("industry"))
        duration = job.get("duration_months")
        duration_str = f"{duration} months" if duration is not None else NOT_PROVIDED
        is_current = job.get("is_current", False)
        current_tag = " (current)" if is_current else ""
        description = _safe_str(job.get("description"))
        parts.append(
            f"{title} at {company} [{industry}] – {duration_str}{current_tag}. "
            f"{description}"
        )
    return " | ".join(parts)


def _format_skills(skills: list[dict]) -> str:
    """Flatten skills list into a weighted text representation."""
    if not skills:
        return NOT_PROVIDED
    parts: list[str] = []
    for s in skills:
        name = _safe_str(s.get("name"))
        proficiency = _safe_str(s.get("proficiency"))
        endorsements = s.get("endorsements", 0)
        duration = s.get("duration_months")
        dur_str = f"{duration}m experience" if duration else ""
        parts.append(
            f"{name} ({proficiency}, {endorsements} endorsements{', ' + dur_str if dur_str else ''})"
        )
    return ", ".join(parts)


def _format_education(education: list[dict]) -> str:
    """Format education entries."""
    if not education:
        return NOT_PROVIDED
    parts: list[str] = []
    for edu in education:
        degree = _safe_str(edu.get("degree"))
        field_of_study = _safe_str(edu.get("field_of_study"))
        institution = _safe_str(edu.get("institution"))
        end_year = _safe_str(edu.get("end_year"))
        grade = _safe_str(edu.get("grade"))
        tier = _safe_str(edu.get("tier"))
        parts.append(
            f"{degree} in {field_of_study} from {institution} "
            f"(graduated {end_year}, grade: {grade}, tier: {tier})"
        )
    return " | ".join(parts)


def _format_certifications(certs: list[dict]) -> str:
    if not certs:
        return NOT_PROVIDED
    return ", ".join(
        f"{c.get('name', NOT_PROVIDED)} by {c.get('issuer', NOT_PROVIDED)} ({c.get('year', NOT_PROVIDED)})"
        for c in certs
    )


def _format_languages(languages: list[dict]) -> str:
    if not languages:
        return NOT_PROVIDED
    return ", ".join(
        f"{l.get('language', NOT_PROVIDED)} ({l.get('proficiency', NOT_PROVIDED)})"
        for l in languages
    )


def _format_signals(signals: dict) -> str:
    """Summarise redrob platform signals as a compact text block."""
    if not signals:
        return NOT_PROVIDED

    open_to_work = "open to work" if signals.get("open_to_work_flag") else "not actively looking"
    relocate = "willing to relocate" if signals.get("willing_to_relocate") else "not willing to relocate"
    work_mode = _safe_str(signals.get("preferred_work_mode"))
    notice = signals.get("notice_period_days")
    notice_str = f"{notice} days notice" if notice is not None else NOT_PROVIDED

    salary = signals.get("expected_salary_range_inr_lpa", {})
    sal_min = salary.get("min", NOT_PROVIDED)
    sal_max = salary.get("max", NOT_PROVIDED)
    sal_str = f"Expected salary: {sal_min}–{sal_max} LPA INR"

    github = signals.get("github_activity_score", -1)
    github_str = f"GitHub activity score: {github}" if github != -1 else "No GitHub linked"

    completeness = signals.get("profile_completeness_score", NOT_PROVIDED)
    response_rate = signals.get("recruiter_response_rate", NOT_PROVIDED)
    interview_rate = signals.get("interview_completion_rate", NOT_PROVIDED)

    assessment_scores = signals.get("skill_assessment_scores", {})
    if assessment_scores:
        assess_str = "Assessment scores: " + ", ".join(
            f"{k}: {v}" for k, v in assessment_scores.items()
        )
    else:
        assess_str = "No assessments taken"

    return (
        f"{open_to_work}; {relocate}; preferred work mode: {work_mode}; "
        f"{notice_str}; {sal_str}; {github_str}; "
        f"Profile completeness: {completeness}%; "
        f"Recruiter response rate: {response_rate}; "
        f"Interview completion rate: {interview_rate}; "
        f"{assess_str}"
    )


# ---------------------------------------------------------------------------
# Core: Build rich profile string
# ---------------------------------------------------------------------------

def build_rich_profile(candidate: dict) -> str:
    """
    Build a comprehensive textual representation of a candidate suitable
    for semantic embedding.

    Parameters
    ----------
    candidate : dict
        A parsed candidate record from candidates.jsonl.

    Returns
    -------
    str
        A rich natural-language summary of the candidate.
    """
    cid = _safe_str(candidate.get("candidate_id"))

    # --- Profile block ---
    profile = candidate.get("profile", {})
    name = _safe_str(profile.get("anonymized_name"))
    headline = _safe_str(profile.get("headline"))
    summary = _safe_str(profile.get("summary"))
    location = _safe_str(profile.get("location"))
    country = _safe_str(profile.get("country"))
    yoe = _safe_str(profile.get("years_of_experience"))
    current_title = _safe_str(profile.get("current_title"))
    current_company = _safe_str(profile.get("current_company"))
    current_industry = _safe_str(profile.get("current_industry"))
    company_size = _safe_str(profile.get("current_company_size"))

    career_text = _format_career_history(candidate.get("career_history", []))
    education_text = _format_education(candidate.get("education", []))
    skills_text = _format_skills(candidate.get("skills", []))
    certs_text = _format_certifications(candidate.get("certifications", []))
    lang_text = _format_languages(candidate.get("languages", []))
    signals_text = _format_signals(candidate.get("redrob_signals", {}))

    rich_profile = (
        f"Candidate ID: {cid}. "
        f"Name: {name}. "
        f"Current Role: {current_title} at {current_company} ({current_industry}, size: {company_size}). "
        f"Headline: {headline}. "
        f"Location: {location}, {country}. "
        f"Years of Experience: {yoe}. "
        f"Professional Summary: {summary} "
        f"Career History: {career_text} "
        f"Education: {education_text} "
        f"Skills: {skills_text} "
        f"Certifications: {certs_text} "
        f"Languages: {lang_text} "
        f"Platform Signals: {signals_text}"
    )
    return rich_profile


# ---------------------------------------------------------------------------
# Data Loading
# ---------------------------------------------------------------------------

def load_candidates(jsonl_path: Path) -> list[CandidateRecord]:
    """
    Load and preprocess all candidates from a JSONL file.

    Parameters
    ----------
    jsonl_path : Path
        Path to the candidates.jsonl file.

    Returns
    -------
    list[CandidateRecord]
        List of parsed and enriched candidate records.
    """
    logger.info("Loading candidates from: %s", jsonl_path)
    records: list[CandidateRecord] = []
    errors = 0

    with open(jsonl_path, "r", encoding="utf-8") as f:
        for line_num, line in enumerate(tqdm(f, desc="Parsing JSONL", unit=" lines"), start=1):
            line = line.strip()
            if not line:
                continue
            try:
                raw = json.loads(line)
                cid = raw.get("candidate_id", f"UNKNOWN_{line_num}")
                rich_profile = build_rich_profile(raw)
                records.append(CandidateRecord(
                    candidate_id=cid,
                    rich_profile=rich_profile,
                    raw=raw,
                ))
            except json.JSONDecodeError as exc:
                logger.warning("JSON parse error on line %d: %s", line_num, exc)
                errors += 1

    logger.info(
        "Loaded %d candidates successfully. Skipped %d malformed lines.",
        len(records), errors
    )
    return records


# ---------------------------------------------------------------------------
# Embedding Engine
# ---------------------------------------------------------------------------

def embed_profiles(
    records: list[CandidateRecord],
    model_name: str = EMBEDDING_MODEL,
    batch_size: int = BATCH_SIZE,
) -> np.ndarray:
    """
    Embed rich_profile strings using a sentence-transformer model.

    Parameters
    ----------
    records : list[CandidateRecord]
        Candidate records whose rich_profile will be embedded.
    model_name : str
        HuggingFace model identifier.
    batch_size : int
        Number of profiles to encode per batch.

    Returns
    -------
    np.ndarray
        Float32 array of shape (N, embedding_dim).
    """
    logger.info("Loading embedding model: %s", model_name)
    model = SentenceTransformer(model_name)

    texts = [r.rich_profile for r in records]
    logger.info("Encoding %d profiles (batch_size=%d)…", len(texts), batch_size)
    t0 = time.time()

    embeddings = model.encode(
        texts,
        batch_size=batch_size,
        show_progress_bar=True,
        convert_to_numpy=True,
        normalize_embeddings=True,   # cosine similarity via inner product
    )

    elapsed = time.time() - t0
    logger.info(
        "Encoding complete: %d vectors, dim=%d, took %.1fs",
        len(embeddings), embeddings.shape[1], elapsed
    )
    return embeddings.astype(np.float32)


# ---------------------------------------------------------------------------
# FAISS Index
# ---------------------------------------------------------------------------

def build_faiss_index(embeddings: np.ndarray) -> faiss.IndexFlatIP:
    """
    Build a FAISS flat inner-product index (equivalent to cosine similarity
    when embeddings are L2-normalised).

    Parameters
    ----------
    embeddings : np.ndarray
        Shape (N, D) float32 matrix of L2-normalised embeddings.

    Returns
    -------
    faiss.IndexFlatIP
        Populated FAISS index.
    """
    dim = embeddings.shape[1]
    logger.info("Building FAISS IndexFlatIP (dim=%d, n=%d)…", dim, len(embeddings))
    index = faiss.IndexFlatIP(dim)
    index.add(embeddings)
    logger.info("FAISS index built. Total vectors: %d", index.ntotal)
    return index


def save_index(
    index: faiss.Index,
    records: list[CandidateRecord],
    index_path: Path,
    metadata_path: Path,
) -> None:
    """
    Persist the FAISS index and candidate metadata to disk.

    Parameters
    ----------
    index : faiss.Index
        Populated FAISS index.
    records : list[CandidateRecord]
        Original candidate records (used for metadata mapping).
    index_path : Path
        Output path for the .faiss index file.
    metadata_path : Path
        Output path for the JSON metadata file.
    """
    faiss.write_index(index, str(index_path))
    logger.info("FAISS index saved → %s", index_path)

    metadata = {
        str(i): {
            "candidate_id": rec.candidate_id,
            "rich_profile": rec.rich_profile,
        }
        for i, rec in enumerate(records)
    }
    with open(metadata_path, "w", encoding="utf-8") as f:
        json.dump(metadata, f, ensure_ascii=False, indent=2)
    logger.info("Metadata saved → %s", metadata_path)


def load_index(
    index_path: Path,
    metadata_path: Path,
) -> tuple[faiss.Index, dict]:
    """
    Load a persisted FAISS index and its metadata from disk.

    Parameters
    ----------
    index_path : Path
        Path to the .faiss index file.
    metadata_path : Path
        Path to the metadata JSON file.

    Returns
    -------
    tuple[faiss.Index, dict]
        (FAISS index, metadata dict keyed by str integer index).
    """
    if not index_path.exists() or not metadata_path.exists():
        raise FileNotFoundError(
            f"Index or metadata not found. "
            f"Run with --rebuild to build them first.\n"
            f"  Expected: {index_path}\n"
            f"           {metadata_path}"
        )
    logger.info("Loading FAISS index from: %s", index_path)
    index = faiss.read_index(str(index_path))
    with open(metadata_path, "r", encoding="utf-8") as f:
        metadata = json.load(f)
    logger.info("Index loaded. Total vectors: %d", index.ntotal)
    return index, metadata


# ---------------------------------------------------------------------------
# Retrieval
# ---------------------------------------------------------------------------

@dataclass
class RetrievalResult:
    """A single candidate retrieval result."""
    rank: int
    candidate_id: str
    similarity_score: float
    rich_profile: str


def get_top_k_candidates(
    job_description_text: str,
    k: int = 50,
    model_name: str = EMBEDDING_MODEL,
    index_path: Path = INDEX_PATH,
    metadata_path: Path = METADATA_PATH,
) -> list[RetrievalResult]:
    """
    Embed a job description and retrieve the top-K most similar candidates
    from the FAISS index.

    Parameters
    ----------
    job_description_text : str
        Raw text of the job description.
    k : int
        Number of top candidates to return.
    model_name : str
        Sentence-transformer model to use for embedding the JD.
    index_path : Path
        Path to the saved FAISS index.
    metadata_path : Path
        Path to the saved metadata JSON.

    Returns
    -------
    list[RetrievalResult]
        Ranked list of candidate results with similarity scores.
    """
    index, metadata = load_index(index_path, metadata_path)
    total = index.ntotal
    k_capped = min(k, total)
    if k_capped < k:
        logger.warning("Requested k=%d but index only has %d vectors; using k=%d.", k, total, k_capped)

    logger.info("Embedding job description (%d chars)…", len(job_description_text))
    model = SentenceTransformer(model_name)
    jd_vec = model.encode(
        [job_description_text],
        convert_to_numpy=True,
        normalize_embeddings=True,
    ).astype(np.float32)

    logger.info("Querying FAISS index for top-%d candidates…", k_capped)
    distances, indices = index.search(jd_vec, k_capped)

    results: list[RetrievalResult] = []
    for rank, (faiss_idx, score) in enumerate(
        zip(indices[0], distances[0]), start=1
    ):
        if faiss_idx == -1:          # FAISS sentinel for "not found"
            continue
        meta = metadata.get(str(faiss_idx), {})
        results.append(RetrievalResult(
            rank=rank,
            candidate_id=meta.get("candidate_id", f"UNKNOWN_{faiss_idx}"),
            similarity_score=float(score),
            rich_profile=meta.get("rich_profile", NOT_PROVIDED),
        ))

    return results


# ---------------------------------------------------------------------------
# JD Extraction
# ---------------------------------------------------------------------------

def extract_jd_text(docx_path: Path) -> str:
    """
    Extract all text from a .docx job description file.

    Parameters
    ----------
    docx_path : Path
        Path to the Word document.

    Returns
    -------
    str
        Full text content of the document.
    """
    logger.info("Extracting job description from: %s", docx_path)
    doc = Document(str(docx_path))
    paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
    # Also pull text from tables (some JDs use table layouts)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text and cell_text not in paragraphs:
                    paragraphs.append(cell_text)
    full_text = "\n".join(paragraphs)
    logger.info("Extracted %d characters from job description.", len(full_text))
    return full_text


# ---------------------------------------------------------------------------
# Build pipeline (index construction)
# ---------------------------------------------------------------------------

def build_pipeline(
    jsonl_path: Path = CANDIDATES_JSONL,
    index_path: Path = INDEX_PATH,
    metadata_path: Path = METADATA_PATH,
) -> tuple[faiss.Index, list[CandidateRecord]]:
    """
    Full build pipeline: load → preprocess → embed → index → save.

    Parameters
    ----------
    jsonl_path : Path
        Source JSONL file.
    index_path : Path
        Destination .faiss index file.
    metadata_path : Path
        Destination metadata JSON file.

    Returns
    -------
    tuple[faiss.Index, list[CandidateRecord]]
        The built index and the candidate records.
    """
    records = load_candidates(jsonl_path)
    if not records:
        raise ValueError("No candidate records loaded. Check the JSONL file.")

    embeddings = embed_profiles(records)
    index = build_faiss_index(embeddings)
    save_index(index, records, index_path, metadata_path)
    return index, records


# ---------------------------------------------------------------------------
# Reporting helpers
# ---------------------------------------------------------------------------

def results_to_dataframe(results: list[RetrievalResult]) -> pd.DataFrame:
    """Convert retrieval results to a pandas DataFrame for easy inspection."""
    return pd.DataFrame([
        {
            "rank": r.rank,
            "candidate_id": r.candidate_id,
            "similarity_score": round(r.similarity_score, 6),
            "profile_snippet": r.rich_profile[:200] + "…",
        }
        for r in results
    ])


def print_top_results(results: list[RetrievalResult], n: int = 10) -> None:
    """Pretty-print the top-N results to stdout."""
    print("\n" + "=" * 72)
    print(f"  TOP {len(results)} CANDIDATES  (showing first {min(n, len(results))})")
    print("=" * 72)
    for r in results[:n]:
        print(f"\n  Rank #{r.rank:>3}  |  {r.candidate_id}  |  Score: {r.similarity_score:.4f}")
        snippet = r.rich_profile[:300].replace("\n", " ")
        print(f"  {snippet}…")
    print("\n" + "=" * 72)


# ---------------------------------------------------------------------------
# CLI entry point
# ---------------------------------------------------------------------------

def _parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Candidate Ranking System – Module 1",
        formatter_class=argparse.RawDescriptionHelpFormatter,
    )
    parser.add_argument(
        "--rebuild",
        action="store_true",
        help="Force rebuild of the FAISS index even if it already exists.",
    )
    parser.add_argument(
        "--candidates", type=Path, default=CANDIDATES_JSONL,
        help="Path to candidates.jsonl (default: %(default)s)",
    )
    parser.add_argument(
        "--jd", type=Path, default=JD_PATH,
        help="Path to job_description.docx (default: %(default)s)",
    )
    parser.add_argument(
        "--index", type=Path, default=INDEX_PATH,
        help="Path to save/load FAISS index (default: %(default)s)",
    )
    parser.add_argument(
        "--metadata", type=Path, default=METADATA_PATH,
        help="Path to save/load metadata JSON (default: %(default)s)",
    )
    parser.add_argument(
        "--top-k", type=int, default=50,
        help="Number of top candidates to retrieve (default: %(default)s)",
    )
    parser.add_argument(
        "--output-csv", type=Path, default=None,
        help="If set, save results DataFrame to this CSV path.",
    )
    return parser.parse_args()


if __name__ == "__main__":
    args = _parse_args()

    # ---- Step 1: Build or load the index ----
    if args.rebuild or not args.index.exists() or not args.metadata.exists():
        logger.info("Building FAISS index from scratch…")
        build_pipeline(
            jsonl_path=args.candidates,
            index_path=args.index,
            metadata_path=args.metadata,
        )
    else:
        logger.info(
            "Index already exists at '%s'. Use --rebuild to regenerate. "
            "Skipping build phase.",
            args.index,
        )

    # ---- Step 2: Extract job description text ----
    jd_text = extract_jd_text(args.jd)

    # ---- Step 3: Retrieve top-K candidates ----
    results = get_top_k_candidates(
        job_description_text=jd_text,
        k=args.top_k,
        index_path=args.index,
        metadata_path=args.metadata,
    )

    # ---- Step 4: Display results ----
    print_top_results(results, n=10)

    df = results_to_dataframe(results)
    logger.info("\nFull results DataFrame (shape %s):\n%s", df.shape, df.to_string(index=False))

    if args.output_csv:
        df_full = pd.DataFrame([
            {
                "rank": r.rank,
                "candidate_id": r.candidate_id,
                "similarity_score": r.similarity_score,
                "rich_profile": r.rich_profile,
            }
            for r in results
        ])
        df_full.to_csv(args.output_csv, index=False)
        logger.info("Results saved to CSV → %s", args.output_csv)

    logger.info("Done.")
